# This file is part of [accountancy-support-tools]. 
# # [accountancy-support-tools] is free software: you can redistribute it and/or 
# # modify it under the terms of the GNU General Public License as published by 
# the Free Software Foundation, either version 3 of the License, or (at your option) 
# any later version. 
# # [accountancy-support-tools] is distributed in the hope that it will be useful, 
# but WITHOUT ANY WARRANTY; without even the implied warranty of 
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE. See the 
# GNU General Public License for more details. 
# # You should have received a copy of the GNU General Public License 
# # along with [accountancy-support-tools]. If not, see <http://www.gnu.org/licenses/>.



from docx import Document
import openpyxl
from openpyxl import Workbook
from openpyxl import load_workbook
import os
import datetime
import random
import shutil
from lxml import etree
from functools import wraps
import sys
from PyQt5 import QtWidgets



app = QtWidgets.QApplication(sys.argv)
def suppress_errors(func):
    error_dialog = QtWidgets.QErrorMessage()
    @wraps(func)
    def wrapper(*args, **kwargs):
        try:
            return func(*args, **kwargs)
        except Exception as e:
            error_dialog.showMessage(f"Error in function '{func.__name__}': {e}")
            error_dialog.exec_()
    return wrapper



@suppress_errors
def replace_text_in_paragraph(paragraph, data, cell=None):
    for key, value in data.items():
        if key in paragraph.text:
            inline = paragraph.runs
            for i in range(len(inline)):
                inline[i].text = inline[i].text.replace(key,str(value))      
        if cell is not None:
            cell.text = cell.text.replace(key, str(data.get(key, "")))



@suppress_errors
def replace_text_in_cell(cell, data):                               
    for paragraph in cell.paragraphs:    
        replace_text_in_paragraph(paragraph,data,cell)   
    for table in cell.tables:                
        replace_text_in_table(table, data)   



@suppress_errors
def replace_text_in_table(table, data):
    for row in table.rows:
        for cell in row.cells:
            replace_text_in_cell(cell, data)



@suppress_errors
def replace_text_in_textbox(textbox, data):
    for paragraph in textbox.paragraphs:
        replace_text_in_paragraph(paragraph, data)



@suppress_errors
def replace_text_in_shape(shape, data):
    if shape.text is not None:
        for key, value in data.items():
            if key in shape.text:
                shape.text = shape.text.replace(key, str(value))



@suppress_errors
def replace_text_in_drawing(drawing, data):
    xml = etree.tostring(drawing, encoding='unicode')
    for key, value in data.items():
        xml = xml.replace(key, str(value))
    new_drawing = etree.fromstring(xml)
    drawing.getparent().replace(drawing, new_drawing)



def inject_data(template_path, output_path, data):
    template = Document(template_path)

    for paragraph in template.paragraphs:
        replace_text_in_paragraph(paragraph, data)

    for table in template.tables:
        replace_text_in_table(table, data)

    for shape in template.inline_shapes:
        print(f'Shape type: {shape.type}')
        if shape.type == 3:
            replace_text_in_textbox(shape, data)

    namespace = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    for drawing in template.element.findall('.//w:drawing', namespace):
        replace_text_in_drawing(drawing, data)               
  
    template.save(output_path)



def generate_word_reports(table_path, template_path, sheet_name):      
    
    filename = os.path.join(f"{table_path}")
    wb = load_workbook(filename, data_only=True)
    ws = wb[sheet_name]
    data = {}
    
    for row in range(3, ws.max_row + 1):
        for col in range(1,ws.max_column + 1):
            key = ws.cell(2, col).value
            value = ws.cell(row, col).value
            if key is not None:
                data[key] = value
            else:
                break
        
        current_datetime = datetime.datetime.now()
        formatted_datetime = current_datetime.strftime("%Y-%m-%d__%H-%M-%S")

        if '{FileName}' in data: 
            output_path = f"output/{data['{FileName}']}__{formatted_datetime}.docx"
        else:
            random_number = str(random.randint(111111, 999999))
            output_path = f"output/{random_number}__{formatted_datetime}.docx"

        inject_data(template_path, output_path, data)



def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)



if __name__ == '__main__':
    template_xlsx_path = resource_path('templates/table.xlsx')
    template_docx_path = resource_path('templates/template.docx')
    readme_path = resource_path('templates/README.md')

    table_path = 'input/table.xlsx'
    sheet_name = 'machine_sheet'
    template_path = 'input/template.docx'

    os.makedirs('input', exist_ok=True)
    os.makedirs('output', exist_ok=True)

    if not os.path.exists('input/table.xlsx'):
        shutil.copy(template_xlsx_path, 'input/table.xlsx')
    if not os.path.exists('input/template.docx'):
        shutil.copy(template_docx_path, 'input/template.docx')
    if not os.path.exists('input/README.md'):
        shutil.copy(readme_path, 'input/README.md')

    generate_word_reports(table_path, template_path, sheet_name)